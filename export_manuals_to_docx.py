from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
MANUALS_DIR = ROOT / "docs" / "manuals"
OUTPUT_DIR = MANUALS_DIR / "docx"
EXTRA_MANUALS = [ROOT / "USER-GUIDE.md"]


def ensure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    if "Manual Bullet" not in styles:
        style = styles.add_style("Manual Bullet", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["List Bullet"]
        style.font.name = "Aptos"
        style.font.size = Pt(11)

    if "Manual Number" not in styles:
        style = styles.add_style("Manual Number", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["List Number"]
        style.font.name = "Aptos"
        style.font.size = Pt(11)


def extract_title_and_subtitle(lines: list[str]) -> tuple[str, str | None]:
    title = "Dashboard Manual"
    subtitle = None
    for idx, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped[2:].strip() or title
            for follow in lines[idx + 1:]:
                follow_stripped = follow.strip()
                if not follow_stripped:
                    continue
                if follow_stripped.startswith(">"):
                    subtitle = follow_stripped[1:].strip()
                break
            break
    return title, subtitle


def _append_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def add_page_number_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.clear()
        paragraph.add_run("Page ")
        _append_field(paragraph, " PAGE ")
        paragraph.add_run(" of ")
        _append_field(paragraph, " NUMPAGES ")


def add_title_page(doc: Document, title: str, subtitle: str | None) -> None:
    doc.add_paragraph("")
    doc.add_paragraph("")
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(24)

    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sub_run = sub.add_run(subtitle)
        sub_run.font.name = "Aptos"
        sub_run.font.size = Pt(12)

    generated = doc.add_paragraph()
    generated.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    gen_run = generated.add_run(f"Generated: {Path.cwd().name}")
    gen_run.font.name = "Aptos"
    gen_run.font.size = Pt(11)

    stamp = doc.add_paragraph()
    stamp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stamp_run = stamp.add_run("Mind Over Mountains Dashboard Guide")
    stamp_run.font.name = "Aptos"
    stamp_run.font.size = Pt(11)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_contents_page(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    paragraph = doc.add_paragraph()
    _append_field(paragraph, r' TOC \o "1-3" \h \z \u ')
    note = doc.add_paragraph()
    note.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    note_run = note.add_run("Update the table in Word if entries do not appear immediately.")
    note_run.italic = True
    note_run.font.size = Pt(10)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def render_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if cells and not all(set(cell) <= {"-", ":"} for cell in cells):
            rows.append(cells)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            table.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""


def add_image(doc: Document, manual_path: Path, rel_path: str) -> None:
    image_path = (manual_path.parent / rel_path).resolve()
    if image_path.suffix.lower() == ".svg":
        png_candidate = image_path.with_suffix(".png")
        if png_candidate.exists():
            image_path = png_candidate
    if not image_path.exists():
        doc.add_paragraph(f"[Missing image: {rel_path}]")
        return
    doc.add_picture(str(image_path), width=Inches(6.8))


def add_paragraph_with_formatting(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def convert_manual(manual_path: Path) -> Path:
    doc = Document()
    ensure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.start_type = WD_SECTION.NEW_PAGE

    lines = manual_path.read_text(encoding="utf-8").splitlines()
    title, subtitle = extract_title_and_subtitle(lines)
    add_title_page(doc, title, subtitle)
    add_contents_page(doc)
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            render_table(doc, table_lines)
            continue

        if stripped.startswith("!["):
            match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", stripped)
            if match:
                add_image(doc, manual_path, match.group(1))
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            doc.add_heading(title, level=min(level - 1, 4))
            i += 1
            continue

        if stripped.startswith(">"):
            add_paragraph_with_formatting(doc, stripped[1:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_paragraph_with_formatting(doc, re.sub(r"^\d+\.\s+", "", stripped), style="Manual Number")
            i += 1
            continue

        if stripped.startswith("- "):
            add_paragraph_with_formatting(doc, stripped[2:], style="Manual Bullet")
            i += 1
            continue

        add_paragraph_with_formatting(doc, stripped)
        i += 1

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    add_page_number_footer(doc)
    output_path = OUTPUT_DIR / f"{manual_path.stem}.docx"
    doc.save(output_path)
    return output_path


def main() -> None:
    manual_paths = sorted(MANUALS_DIR.glob("*.md"))
    manual_paths.extend([path for path in EXTRA_MANUALS if path.exists()])
    for manual_path in manual_paths:
        output_path = convert_manual(manual_path)
        print(output_path)


if __name__ == "__main__":
    main()
